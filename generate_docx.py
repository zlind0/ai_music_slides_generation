#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
通用教案生成器
用法:
  python3 generate_docx.py <data.json>        # 生成单个教案
  python3 generate_docx.py <folder/>          # 批量生成文件夹下所有 .json

数据文件格式见 lessons/README.json
"""
import json
import sys
import os
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_run_font(run, latin="Times New Roman", east_asia="宋体"):
    """同时设置西文字体（Times New Roman）和中文字体（宋体）。"""
    run.font.name = latin
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.insert(0, rFonts)
    rFonts.set(qn("w:eastAsia"), east_asia)
    rFonts.set(qn("w:ascii"), latin)
    rFonts.set(qn("w:hAnsi"), latin)


def set_cell_text(cell, text, bold=False, font_size=11, align=None):
    cell.text = ""
    para = cell.paragraphs[0]
    if align:
        para.alignment = align
    run = para.add_run(text)
    run.bold = bold
    run.font.size = Pt(font_size)
    set_run_font(run)


def create_lesson_plan(data: dict, output_path: str):
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # 顶部标题
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run("德兴市二中（音乐 ）电子集体备课")
    title_run.bold = True
    title_run.font.size = Pt(14)
    set_run_font(title_run)

    table = doc.add_table(rows=0, cols=9)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Row 0: 基本信息
    r = table.add_row()
    r.cells[0].merge(r.cells[7])
    set_cell_text(r.cells[0], "基本信息", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(r.cells[8], "")

    # Row 1: 年级/学科/教师
    r = table.add_row()
    r.cells[0].merge(r.cells[1])
    set_cell_text(r.cells[0], "年级", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(r.cells[2], "八", align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(r.cells[3], "学科", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(r.cells[4], "音乐", align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(r.cells[5], "主备教师", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(r.cells[6], "林辰钰", align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(r.cells[7], "附备教师", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(r.cells[8], "")

    # Row 2: 课题
    r = table.add_row()
    r.cells[0].merge(r.cells[1])
    set_cell_text(r.cells[0], "课题", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    r.cells[2].merge(r.cells[7])
    set_cell_text(r.cells[2], data["title"], align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(r.cells[8], "")

    # Row 3: 素养目标 header
    r = table.add_row()
    r.cells[0].merge(r.cells[7])
    set_cell_text(r.cells[0], "素养目标", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(r.cells[8], "二次备课", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)

    # Row 4: 三维目标内容
    r = table.add_row()
    r.cells[0].merge(r.cells[7])
    goals = (
        "知识目标：" + data["knowledge_goal"] + "\n"
        + "能力目标：" + data["ability_goal"] + "\n"
        + "情感目标：" + data["emotion_goal"]
    )
    set_cell_text(r.cells[0], goals)
    set_cell_text(r.cells[8], "")

    # Row 5: 教学重点和难点 header
    r = table.add_row()
    r.cells[0].merge(r.cells[7])
    set_cell_text(r.cells[0], "教学重点和难点", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(r.cells[8], "")

    # Row 6: 重点/难点内容
    r = table.add_row()
    r.cells[0].merge(r.cells[7])
    set_cell_text(r.cells[0], "重点：" + data["focus"] + "\n难点：" + data["difficulty"])
    set_cell_text(r.cells[8], "")

    # Row 7: 教学过程 header
    r = table.add_row()
    r.cells[0].merge(r.cells[7])
    set_cell_text(r.cells[0], "教学过程", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(r.cells[8], "")

    # Row 8: 教学环节 + 导入
    r = table.add_row()
    set_cell_text(r.cells[0], "教学环节", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    r.cells[1].merge(r.cells[7])
    set_cell_text(r.cells[1], data["intro"])
    set_cell_text(r.cells[8], "")

    # Row 9: 创设情境/引入课题
    r = table.add_row()
    set_cell_text(r.cells[0], "创设情境\n\n引入课题", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    r.cells[1].merge(r.cells[7])
    set_cell_text(r.cells[1], data["new_lesson"])
    set_cell_text(r.cells[8], "")

    # Row 10: 探究新知/讲授新课
    r = table.add_row()
    set_cell_text(r.cells[0], "探究新知\n\n\n\n\n\n讲授新课", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    r.cells[1].merge(r.cells[7])
    set_cell_text(r.cells[1], data["consolidate"])
    set_cell_text(r.cells[8], "")

    # Row 11: 巩固应用
    r = table.add_row()
    set_cell_text(r.cells[0], "巩固应用", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    r.cells[1].merge(r.cells[7])
    set_cell_text(r.cells[1], data["summary"])
    set_cell_text(r.cells[8], "")

    # Row 12: 四小结
    r = table.add_row()
    set_cell_text(r.cells[0], "四小结", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    r.cells[1].merge(r.cells[7])
    set_cell_text(r.cells[1], "总结今天所学内容。")
    set_cell_text(r.cells[8], "")

    # Row 13: 板书设计
    r = table.add_row()
    r.cells[0].merge(r.cells[8])
    set_cell_text(r.cells[0], data["board"])

    doc.save(output_path)
    print("已生成：" + output_path)


def process_json(json_path: str):
    with open(json_path, encoding="utf-8") as f:
        data = json.load(f)

    # 输出路径：同目录下同名 .docx，或 data 里指定 output_path
    if "output_path" in data:
        out = data["output_path"]
    else:
        base = os.path.splitext(json_path)[0]
        out = base + ".docx"

    create_lesson_plan(data, out)


def main():
    if len(sys.argv) < 2:
        print("用法: python3 generate_docx.py <data.json|folder/>")
        sys.exit(1)

    target = sys.argv[1]

    if os.path.isdir(target):
        json_files = sorted(
            os.path.join(root, f)
            for root, _, files in os.walk(target)
            for f in files if f.endswith(".json")
        )
        if not json_files:
            print("未找到 .json 文件")
            sys.exit(1)
        for jf in json_files:
            process_json(jf)
    elif os.path.isfile(target):
        process_json(target)
    else:
        print("找不到文件或文件夹：" + target)
        sys.exit(1)


if __name__ == "__main__":
    main()
