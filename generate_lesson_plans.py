#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Generate 八年级下 第二单元 lesson plan docx files,
mimicking the format of 八年级上 examples.
"""
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_cell_text(cell, text, bold=False, font_size=11, align=None):
    cell.text = ""
    para = cell.paragraphs[0]
    if align:
        para.alignment = align
    run = para.add_run(text)
    run.bold = bold
    run.font.size = Pt(font_size)


def create_lesson_plan(output_path, title, knowledge_goal, ability_goal, emotion_goal,
                       focus, difficulty, intro, new_lesson, consolidate, summary, board):
    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run("德兴市二中（音乐 ）电子集体备课")
    title_run.bold = True
    title_run.font.size = Pt(14)

    table = doc.add_table(rows=0, cols=9)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Row 0: 基本信息
    row0 = table.add_row()
    row0.cells[0].merge(row0.cells[7])
    set_cell_text(row0.cells[0], "基本信息", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(row0.cells[8], "")

    # Row 1: 年级/学科
    row1 = table.add_row()
    row1.cells[0].merge(row1.cells[1])
    set_cell_text(row1.cells[0], "年级", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(row1.cells[2], "八", align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(row1.cells[3], "学科", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(row1.cells[4], "音乐", align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(row1.cells[5], "主备教师", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(row1.cells[6], "林辰钰", align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(row1.cells[7], "附备教师", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(row1.cells[8], "")

    # Row 2: 课题
    row2 = table.add_row()
    row2.cells[0].merge(row2.cells[1])
    set_cell_text(row2.cells[0], "课题", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    row2.cells[2].merge(row2.cells[7])
    set_cell_text(row2.cells[2], title, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(row2.cells[8], "")

    # Row 3: 素养目标 header
    row3 = table.add_row()
    row3.cells[0].merge(row3.cells[7])
    set_cell_text(row3.cells[0], "素养目标", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(row3.cells[8], "二次备课", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)

    # Row 4: Goals
    row4 = table.add_row()
    row4.cells[0].merge(row4.cells[7])
    goals_text = "知识目标：" + knowledge_goal + "\n能力目标：" + ability_goal + "\n情感目标：" + emotion_goal
    set_cell_text(row4.cells[0], goals_text)
    set_cell_text(row4.cells[8], "")

    # Row 5: 教学重点和难点 header
    row5 = table.add_row()
    row5.cells[0].merge(row5.cells[7])
    set_cell_text(row5.cells[0], "教学重点和难点", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(row5.cells[8], "")

    # Row 6: Focus/Difficulty
    row6 = table.add_row()
    row6.cells[0].merge(row6.cells[7])
    set_cell_text(row6.cells[0], "重点：" + focus + "\n难点：" + difficulty)
    set_cell_text(row6.cells[8], "")

    # Row 7: 教学过程 header
    row7 = table.add_row()
    row7.cells[0].merge(row7.cells[7])
    set_cell_text(row7.cells[0], "教学过程", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(row7.cells[8], "")

    # Row 8: 导入
    row8 = table.add_row()
    set_cell_text(row8.cells[0], "教学环节", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    row8.cells[1].merge(row8.cells[7])
    set_cell_text(row8.cells[1], intro)
    set_cell_text(row8.cells[8], "")

    # Row 9: 创设情境/新课教授
    row9 = table.add_row()
    set_cell_text(row9.cells[0], "创设情境\n\n引入课题", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    row9.cells[1].merge(row9.cells[7])
    set_cell_text(row9.cells[1], new_lesson)
    set_cell_text(row9.cells[8], "")

    # Row 10: 探究新知/讲授新课
    row10 = table.add_row()
    set_cell_text(row10.cells[0], "探究新知\n\n\n\n\n\n讲授新课", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    row10.cells[1].merge(row10.cells[7])
    set_cell_text(row10.cells[1], consolidate)
    set_cell_text(row10.cells[8], "")

    # Row 11: 巩固应用
    row11 = table.add_row()
    set_cell_text(row11.cells[0], "巩固应用", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    row11.cells[1].merge(row11.cells[7])
    set_cell_text(row11.cells[1], summary)
    set_cell_text(row11.cells[8], "")

    # Row 12: 四小结
    row12 = table.add_row()
    set_cell_text(row12.cells[0], "四小结", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    row12.cells[1].merge(row12.cells[7])
    set_cell_text(row12.cells[1], "总结今天所学内容。")
    set_cell_text(row12.cells[8], "")

    # Row 13: 板书设计
    row13 = table.add_row()
    row13.cells[0].merge(row13.cells[8])
    set_cell_text(row13.cells[0], board)

    doc.save(output_path)
    print("已生成：" + output_path)


# ============================================================
# Song 1: 《但愿人长久》
# ============================================================
S1_knowledge = (
    "了解《但愿人长久》的词曲作者及创作背景，认识宋词与流行音乐融合的艺术特点；"
    "了解作曲家梁弘志及演唱者邓丽君在华语流行音乐中的地位。"
)
S1_ability = (
    "能够用自然、流畅的声音完整演唱《但愿人长久》，准确把握歌曲的旋律走向与情感层次，"
    "学会气息连贯的演唱技巧。"
)
S1_emotion = (
    "通过演唱感受苏轼词作中旷达豁然的人生境界，体会『但愿人长久，千里共婵娟』"
    "所表达的对亲人的美好祝愿和对人生的乐观态度。"
)
S1_focus = (
    "准确演唱歌曲旋律，掌握歌曲的节奏特点（如切分节奏）及情感处理；理解词曲结合的艺术特色。"
)
S1_difficulty = (
    "表现歌曲中从『问天』的惆怅到『但愿人长久』的豁然开朗的情感转变；"
    "保持气息连贯，避免断句影响旋律流畅性。"
)
S1_intro = (
    "（一）导入\n"
    "1. 情境创设：\n"
    "    播放中秋节相关图片或视频（月圆、灯笼、家人团聚），提问：\n"
    "    『中秋节的到来让你想到哪些关于月亮的诗词？』\n"
    "    学生自由回答（如『举头望明月』、『海上生明月』等），教师引导：\n"
    "    『今天我们要学习的歌曲，就是根据宋代大文豪苏轼一首著名的词谱曲而成的。』\n"
    "    板书：苏轼《水调歌头》——『但愿人长久，千里共婵娟』\n"
    "    朗读歌词全文，体会词作情感。"
)
S1_new_lesson = (
    "（二）新课教授\n"
    "1. 初听歌曲，整体感知\n"
    "    播放邓丽君演唱版《但愿人长久》，学生聆听，思考：\n"
    "      ① 歌曲的情绪是什么？（婉转深情、略带忧思而最终豁然）\n"
    "      ② 歌曲节拍与速度特点？（中速，旋律流畅舒展）\n"
    "2. 背景讲解\n"
    "    苏轼于1076年中秋，因思念弟弟苏辙写下《水调歌头》，抒发对亲人的思念，\n"
    "    最终以『但愿人长久，千里共婵娟』升华，表达旷达的人生态度。\n"
    "    作曲家梁弘志（1954—2004）：台湾著名词曲作家，曾创作《恰似你的温柔》等经典作品。\n"
    "    演唱者邓丽君：华语流行音乐的代表人物，以细腻婉转的演唱风格著称。"
)
S1_consolidate = (
    "3. 学唱歌曲\n"
    "    分段教唱：\n"
    "      ① 先随音乐哼唱旋律（『lu』），感受旋律线条的起伏。\n"
    "      ② 填词跟唱，注意换气点（每句末尾）。\n"
    "      ③ 难点突破：\n"
    "         - 『转朱阁，低绮户，照无眠』的切分节奏，强调节奏准确性；\n"
    "         - 『但愿人长久』一句情绪渐强，气息要支撑住；\n"
    "         - 结尾『千里共婵娟』长音，气息平稳渐弱，余音袅袅。\n"
    "    情感处理：\n"
    "      第一段（问天）：声音轻柔，带有疑惑与惆怅；\n"
    "      第二段（转折）：情绪稍紧张，『不应有恨』稍做强调；\n"
    "      结尾（升华）：声音渐明朗，饱含深情与祝福。\n"
    "（三）拓展比较\n"
    "1. 音乐版本对比\n"
    "    播放陆在易版《水调歌头·明月几时有》片段，与梁弘志版对比：\n"
    "      - 陆在易版：学院派风格，旋律更加传统古典，情感含蓄深沉；\n"
    "      - 梁弘志版：流行音乐风格，旋律亲切易唱，情感表达更为直接。\n"
    "    引导学生思考：古诗词谱曲时，流行化处理对传播有何影响？"
)
S1_summary = (
    "播放完整伴奏，指挥学生完整演唱《但愿人长久》。\n"
    "要求：声音自然流畅，情感真实投入，注意气息连贯。\n"
    "选取2—3名同学或小组展示演唱，全班给予评价与鼓励。"
)
S1_board = (
    "板书设计  《但愿人长久》\n"
    "词：[宋]苏轼《水调歌头》    曲：梁弘志    唱：邓丽君\n"
    "1. 情感变化：问天惆怅 → 思人忧愁 → 旷达祝愿\n"
    "2. 音乐特点：中速、流畅旋律、切分节奏、气息连贯\n"
    "3. 文化内涵：诗词与音乐的完美融合，中秋思念的文化意象"
)

create_lesson_plan(
    output_path="/Users/lin/Documents/Code/music_slides/八年级下/第二单元/《但愿人长久》.docx",
    title="《但愿人长久》",
    knowledge_goal=S1_knowledge,
    ability_goal=S1_ability,
    emotion_goal=S1_emotion,
    focus=S1_focus,
    difficulty=S1_difficulty,
    intro=S1_intro,
    new_lesson=S1_new_lesson,
    consolidate=S1_consolidate,
    summary=S1_summary,
    board=S1_board
)

# ============================================================
# Song 2: 《沁园春·雪》
# ============================================================
S2_knowledge = (
    "了解歌曲《沁园春·雪》的创作背景，认识毛泽东诗词的豪迈风格；"
    "了解复二部曲式结构及说唱音乐、戏曲音乐在歌曲中的运用。"
)
S2_ability = (
    "通过欣赏和分析，感受歌曲壮美雄浑的意境，能够识别歌曲速度、力度、旋律的变化，"
    "并能用语言描述音乐形象。"
)
S2_emotion = (
    "感受毛泽东诗词的壮志豪情与革命乐观主义精神，激发学生对祖国山河的热爱之情和文化自信。"
)
S2_focus = (
    "欣赏并分析《沁园春·雪》的音乐特点，包括速度变化、旋律起伏与诗词情感的关系；"
    "理解复二部曲式结构。"
)
S2_difficulty = (
    "理解说唱音乐与戏曲润腔的创作手法在歌曲中的体现；"
    "体会不同段落音乐情绪的对比与转换。"
)
S2_intro = (
    "（一）导入\n"
    "1. 情境创设：\n"
    "    展示1936年北国雪景图片（长城、黄河、茫茫白雪），播放背景音乐。\n"
    "    提问：『如果你站在冰封万里的北国，看到这样的壮景，会有怎样的感受？』\n"
    "    引入毛泽东1936年写下《沁园春·雪》的历史背景：长征胜利后，红军初到陕北，\n"
    "    登高远望北国雪景，写下这首气势磅礴的词作。\n"
    "    引导学生朗读：『北国风光，千里冰封，万里雪飘……』，体会意境。"
)
S2_new_lesson = (
    "（二）欣赏感知\n"
    "1. 初听全曲\n"
    "    完整聆听歌曲（可选张也版或彭丽媛版），感受整体情感。\n"
    "    学生闭眼聆听，思考：歌曲情绪如何变化？哪里舒缓，哪里激昂？\n"
    "2. 介绍作品背景\n"
    "    创作时间：1936年，毛泽东于陕北清涧县袁家沟初稿，1945年在重庆谈判期间发表，震惊文坛。\n"
    "    作曲：生茂、唐诃，将诗词语言音韵融入曲调，旋律随情感变化而委婉起伏。"
)
S2_consolidate = (
    "3. 分段精听与分析\n"
    "    第一段（描景）：\n"
    "      『北国风光，千里冰封，万里雪飘』——舒缓宽广，描绘辽阔北国雪景；\n"
    "      『望长城内外……山舞银蛇，原驰蜡象』——音乐逐渐舒展，如身临其境。\n"
    "    第二段（评古）：\n"
    "      『江山如此多娇，引无数英雄竞折腰』——速度稍快，音调高昂，带有评价气势；\n"
    "      『惜秦皇汉武……略输文采』——说唱性强，吸收戏曲润腔，旋律跌宕起伏。\n"
    "    第三段（展望）：\n"
    "      『数风流人物，还看今朝』——音乐推向高潮，饱满有力，充满革命自信。\n"
    "4. 曲式结构分析\n"
    "    复二部曲式：A段（描景）+ B段（评古展望）\n"
    "    引导学生对照乐谱，标注速度、力度记号的变化。\n"
    "5. 创作手法点拨\n"
    "    旋律以诗词声韵为基础，语调高低直接影响旋律走向；\n"
    "    吸收中国说唱音乐（如大鼓）节奏特点，增强诗词的朗诵感；\n"
    "    融入戏曲润腔（如滑音），体现民族音乐色彩。"
)
S2_summary = (
    "随音乐诵读《沁园春·雪》全词，结合音乐情绪变化有感情地朗诵。\n"
    "小组讨论：毛泽东诗词谱曲后，音乐为诗词增添了什么？（情感渲染、画面感等）\n"
    "推荐欣赏：其他毛泽东诗词歌曲，如《蝶恋花·答李淑一》。"
)
S2_board = (
    "板书设计  《沁园春·雪》\n"
    "词：毛泽东（1936年）    曲：生茂、唐诃\n"
    "曲式：复二部曲式  A（描景·舒缓宽广）+ B（评古展望·激昂高亢）\n"
    "创作手法：诗词声韵入曲 + 说唱节奏 + 戏曲润腔\n"
    "情感：壮志豪情、革命乐观主义"
)

create_lesson_plan(
    output_path="/Users/lin/Documents/Code/music_slides/八年级下/第二单元/《沁园春·雪》.docx",
    title="《沁园春·雪》",
    knowledge_goal=S2_knowledge,
    ability_goal=S2_ability,
    emotion_goal=S2_emotion,
    focus=S2_focus,
    difficulty=S2_difficulty,
    intro=S2_intro,
    new_lesson=S2_new_lesson,
    consolidate=S2_consolidate,
    summary=S2_summary,
    board=S2_board
)

# ============================================================
# Song 3: 《蜀道难》
# ============================================================
S3_knowledge = (
    "了解合唱交响曲《蜀道难》的创作背景及作曲家郭文景；"
    "认识现代作曲技法与四川传统音乐相结合的创作手法；"
    "了解男高音独唱、合唱与管弦乐队的演出形式。"
)
S3_ability = (
    "通过聆听分析，能够感受和辨别作品中不同乐段的情绪变化，"
    "理解音乐如何诠释李白诗歌中『蜀道之难』的意境。"
)
S3_emotion = (
    "感受中国现代音乐创作的探索精神与艺术价值，体会李白诗歌豪放雄奇的风格，"
    "增强对中国传统文化与现代音乐创作的热爱。"
)
S3_focus = (
    "欣赏并感受《蜀道难》的整体音乐形象：险峻、雄奇、神秘；"
    "理解合唱与管弦乐队的配合关系。"
)
S3_difficulty = (
    "理解作品中现代作曲技法（如不协和音响、复杂节奏）与四川传统音乐元素的融合方式；"
    "将音乐音响与诗词意象对应起来。"
)
S3_intro = (
    "（一）导入\n"
    "1. 情境创设：\n"
    "    展示四川蜀道（剑门关、金牛道等）的图片或视频，感受古蜀道的险峻壮观。\n"
    "    提问：『如果要你用音乐来表现这种险峻，你会用什么样的声音？』\n"
    "    朗读李白《蜀道难》开篇：『噫吁嚱，危乎高哉！蜀道之难，难于上青天！』\n"
    "    感受诗歌中夸张而豪放的气魄，引入课题。"
)
S3_new_lesson = (
    "（二）作品介绍\n"
    "1. 创作背景\n"
    "    《蜀道难》是作曲家郭文景根据唐代诗人李白同名诗创作的合唱交响曲。\n"
    "    郭文景（1956年生）：中国当代著名作曲家，作品以中国传统题材为基础，\n"
    "    融合现代西方作曲技法，具有强烈的中国风格。\n"
    "    演出形式：男高音独唱 + 合唱 + 管弦乐队。\n"
    "2. 演奏形式介绍\n"
    "    讲解管弦乐队的组成（弦乐组、木管组、铜管组、打击乐组）；\n"
    "    介绍男高音在歌剧与声乐中的特点：音色明亮、穿透力强，适合表现雄奇豪迈的形象。"
)
S3_consolidate = (
    "3. 欣赏与分析\n"
    "    第一次完整聆听：\n"
    "      学生跟随歌词聆听，思考：哪些段落令你感受到『险』？音乐用了什么手段？\n"
    "    分段精听：\n"
    "      ① 开篇：『噫吁嚱，危乎高哉！』——管弦乐强奏，男高音激昂呐喊，突显惊叹之情；\n"
    "      ② 『黄鹤之飞尚不得过，猿猱欲度愁攀援』——旋律迂回曲折，弦乐颤音模拟山路崎岖；\n"
    "      ③ 『问君西游何时还？』——旋律转为抒情，表现旅途艰辛中的人文关怀；\n"
    "      ④ 『又闻子规啼夜月，愁空山』——弱奏，子规（杜鹃）鸟鸣音效，渲染凄凉意境。\n"
    "4. 创作手法点拨\n"
    "    现代技法：不协和音程制造紧张感，模拟山势险峻；复杂节奏型增强戏剧性冲突；\n"
    "    四川传统音乐：吸收川剧音乐元素（高腔），旋律带有地方色彩，贴近蜀道意境；\n"
    "    对比手法：独唱（个人情感）与合唱（宏大叙事）交替，形成张力。"
)
S3_summary = (
    "再次聆听全曲，鼓励学生闭眼想象画面，聆听后用一段话描述音乐带给你的意境。\n"
    "小组分享：你觉得音乐最成功地表现了诗中哪个意象？\n"
    "拓展思考：与课本《沁园春·雪》对比，两首作品如何用不同方式处理古典诗词？"
)
S3_board = (
    "板书设计  《蜀道难》\n"
    "词：[唐]李白    曲：郭文景\n"
    "演出形式：男高音独唱 + 合唱 + 管弦乐队\n"
    "创作手法：现代作曲技法 + 四川传统音乐（川剧高腔）\n"
    "音乐形象：险峻 → 崎岖 → 抒情 → 凄凉"
)

create_lesson_plan(
    output_path="/Users/lin/Documents/Code/music_slides/八年级下/第二单元/《蜀道难》.docx",
    title="《蜀道难》",
    knowledge_goal=S3_knowledge,
    ability_goal=S3_ability,
    emotion_goal=S3_emotion,
    focus=S3_focus,
    difficulty=S3_difficulty,
    intro=S3_intro,
    new_lesson=S3_new_lesson,
    consolidate=S3_consolidate,
    summary=S3_summary,
    board=S3_board
)

# ============================================================
# Song 4: 《秋之歌》（绝句三首）
# ============================================================
S4_knowledge = (
    "了解《秋之歌》的创作背景及作曲家罗忠镕；"
    "认识三首作品各自的调式、速度、情绪特点；"
    "了解民族化和声与古琴音乐写作手法。"
)
S4_ability = (
    "通过聆听三首不同风格的作品，能够感知和比较各作品的音乐要素（速度、力度、音色）"
    "与情感表达之间的联系，提升音乐感知能力。"
)
S4_emotion = (
    "感受唐诗中浓厚的秋日意境与古典文人情怀，体会中国古诗词音乐化后的独特魅力，"
    "增强对中国古典文化的认同感。"
)
S4_focus = (
    "欣赏《山行》《南陵道中》《寄扬州韩绰判官》三首作品，分析各作品的情感基调与音乐特点。"
)
S4_difficulty = (
    "理解三首作品在音乐风格上的差异（民族化和声 vs 古琴写法），并能用音乐语言加以描述。"
)
S4_intro = (
    "（一）导入\n"
    "1. 情境创设：\n"
    "    展示三幅秋景图——红叶满山、江南水乡秋雨、月下扬州古城，与三首诗对应。\n"
    "    教师有感情地朗读三首诗（《山行》《南陵道中》《寄扬州韩绰判官》），学生感受诗中的意境。\n"
    "    提问：『三首诗同是写秋天，但情感一样吗？』\n"
    "    （《山行》：清新明丽；《南陵道中》：孤独感伤；《寄扬州韩绰判官》：洒脱潇洒）\n"
    "    引入课题：作曲家罗忠镕把这三首诗谱成了一组声乐套曲《秋之歌》。"
)
S4_new_lesson = (
    "（二）作品介绍\n"
    "1. 作曲家介绍\n"
    "    罗忠镕（1924—2019）：中国当代著名作曲家、音乐理论家，长期致力于探索\n"
    "    中国民族音调与现代作曲技法的融合，代表作有《涉江采芙蓉》《秋之歌》等。\n"
    "2. 整体结构\n"
    "    《秋之歌》由三首作品组成（三联歌）：\n"
    "      Ⅰ. 《山行》——[唐]杜牧，1=F，清澈的秋山红叶\n"
    "      Ⅱ. 《南陵道中》——[唐]杜牧，1=bE，感叹秋至的孤客情怀\n"
    "      Ⅲ. 《寄扬州韩绰判官》——1=G，洒脱潇洒的江南秋意"
)
S4_consolidate = (
    "3. 分曲欣赏与分析\n"
    "  Ⅰ.《山行》\n"
    "    聆听，感受速度：徐缓、清澈（=64）；\n"
    "    旋律清澈舒展，描绘『远上寒山石径斜，白云生处有人家』的静谧之美；\n"
    "    『停车坐爱枫林晚，霜叶红于二月花』——情绪转为欣喜，旋律上扬；\n"
    "    分析：民族化和声（五声音阶为基础），体现中国山水画意境。\n"
    "  Ⅱ.《南陵道中》\n"
    "    聆听，感受情绪：感叹、孤寂（=48）；\n"
    "    旋律舒缓低沉，力度p，描绘旅人在秋风中独行的孤独；\n"
    "    『正是客心孤迥处，谁家红袖凭江楼』——一抹温暖色彩，音乐稍起伏；\n"
    "    分析：与第一首对比，速度更慢，情绪更为沉郁。\n"
    "  Ⅲ.《寄扬州韩绰判官》\n"
    "    聆听，感受速度：中速稍快、潇洒（=78）；\n"
    "    旋律活泼灵动，古琴写作风格（仿指法、余音感）；\n"
    "    整体基调轻松，带有好友间的调侃趣味；\n"
    "    分析：借鉴古代琴歌写作手法，体现中国文人音乐的雅趣。\n"
    "4. 综合比较\n"
    "    三首速度对比：快（山行）→ 慢（南陵道中）→ 中（寄扬州）\n"
    "    三种情感：欣喜清新 / 孤独感伤 / 洒脱潇洒"
)
S4_summary = (
    "随音乐配乐朗诵三首唐诗，选择最喜欢的一首进行有感情的演绎。\n"
    "思考题：为什么作曲家把这三首诗组合在一起？它们在音乐上是如何统一的？\n"
    "拓展：推荐欣赏其他古诗词歌曲，如《枫桥夜泊》《春晓》等。"
)
S4_board = (
    "板书设计  《秋之歌》（绝句三首）\n"
    "词：[唐]杜牧    曲：罗忠镕\n"
    "Ⅰ.《山行》  清澈明丽  民族化和声  五声音阶\n"
    "Ⅱ.《南陵道中》  孤独感伤  舒缓低沉  p力度\n"
    "Ⅲ.《寄扬州韩绰判官》  洒脱潇洒  古琴写法  中速活泼\n"
    "共同主题：秋日意境 + 文人情怀"
)

create_lesson_plan(
    output_path="/Users/lin/Documents/Code/music_slides/八年级下/第二单元/《秋之歌》.docx",
    title="《秋之歌》（绝句三首）",
    knowledge_goal=S4_knowledge,
    ability_goal=S4_ability,
    emotion_goal=S4_emotion,
    focus=S4_focus,
    difficulty=S4_difficulty,
    intro=S4_intro,
    new_lesson=S4_new_lesson,
    consolidate=S4_consolidate,
    summary=S4_summary,
    board=S4_board
)

print("\n所有教案已生成完毕！")
